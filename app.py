"""
GamePoch Content Copilot - 游戏发行内容工作台
"""

import asyncio
import json
import os
import re
from datetime import datetime, timedelta
from pathlib import Path
from io import BytesIO

import httpx
from bs4 import BeautifulSoup
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

app = FastAPI(title="GamePoch Content Copilot", version="4.0.0")
BASE_DIR = Path(__file__).parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

# ============================================================
# Steam API
# ============================================================

STEAM = "https://store.steampowered.com/api"

def clean_item(item):
    """清洗Steam API返回的游戏数据"""
    app_id = item.get("id")
    return {
        "name": item.get("name", "未知"),
        "id": app_id,
        "discount_percent": item.get("discount_percent", 0),
        "original_price": item.get("original_price", 0) / 100 if item.get("original_price") else 0,
        "final_price": item.get("final_price", 0) / 100 if item.get("final_price") else 0,
        "header_image": item.get("header_image", ""),
        "url": f"https://store.steampowered.com/app/{app_id}" if app_id else "",
    }


async def get_featured():
    """获取Steam推荐数据（热销/新品/特惠）"""
    async with httpx.AsyncClient(timeout=20) as c:
        r = await c.get(f"{STEAM}/featuredcategories")
        if r.status_code != 200:
            return {"top_sellers": [], "new_releases": [], "specials": []}
        d = r.json()
        return {
            "top_sellers": [clean_item(i) for i in d.get("top_sellers", {}).get("items", [])],
            "new_releases": [clean_item(i) for i in d.get("new_releases", {}).get("items", [])],
            "specials": [clean_item(i) for i in d.get("specials", {}).get("items", [])],
        }


async def search_games(keyword):
    """按名字搜索Steam游戏"""
    async with httpx.AsyncClient(timeout=15) as c:
        r = await c.get(f"{STEAM}/storesearch/", params={"term": keyword, "l": "schinese", "cc": "cn"})
        if r.status_code != 200:
            return []
        try:
            data = r.json()
        except:
            # 处理编码问题
            data = json.loads(r.content.decode('utf-8', errors='replace'))
        items = data.get("items", [])[:8]
        result = []
        for i in items:
            try:
                name = i.get("name", "").encode('utf-8', errors='replace').decode('utf-8')
            except:
                name = str(i.get("id", ""))
            result.append({
                "app_id": i.get("id", 0),
                "name": name,
                "image": i.get("tiny_image", "")
            })
        return result


async def game_details(app_id):
    """获取游戏详情"""
    async with httpx.AsyncClient(timeout=15) as c:
        r = await c.get(f"{STEAM}/appdetails", params={"appids": app_id, "l": "schinese"})
        if r.status_code != 200:
            return None
        d = r.json().get(str(app_id), {})
        if not d.get("success"):
            return None
        info = d["data"]
        return {
            "app_id": app_id,
            "name": info.get("name", ""),
            "description": info.get("short_description", ""),
            "genres": [g["description"] for g in info.get("genres", [])],
            "developers": info.get("developers", []),
            "release_date": info.get("release_date", {}).get("date", ""),
            "header_image": info.get("header_image", ""),
            "url": f"https://store.steampowered.com/app/{app_id}",
        }


async def game_reviews(app_id, num=20):
    """获取游戏评价"""
    async with httpx.AsyncClient(timeout=15) as c:
        r = await c.get(f"https://store.steampowered.com/appreviews/{app_id}", params={
            "json": "1", "filter": "recent", "language": "all", "num_per_page": num
        })
        if r.status_code != 200:
            return {"positive": 0, "negative": 0, "score": "", "reviews": []}
        d = r.json()
        reviews = []
        for rv in d.get("reviews", []):
            reviews.append({
                "up": rv.get("voted_up", True),
                "hours": round(rv.get("author", {}).get("playtime_forever", 0) / 60, 1),
                "text": rv.get("review", ""),
                "lang": rv.get("language", ""),
            })
        s = d.get("query_summary", {})
        return {
            "positive": s.get("total_positive", 0),
            "negative": s.get("total_negative", 0),
            "score": s.get("review_score_desc", ""),
            "reviews": reviews,
        }


# ============================================================
# 翻译
# ============================================================

async def translate(text):
    if not text or len(text.strip()) < 2:
        return text
    async with httpx.AsyncClient(timeout=10) as c:
        r = await c.get("https://translate.googleapis.com/translate_a/single",
                         params={"client": "gtx", "sl": "auto", "tl": "zh-CN", "dt": "t", "q": text[:3000]})
        if r.status_code == 200:
            parts = r.json()[0]
            return "".join(s[0] for s in parts if s[0])
    return text


# ============================================================
# 新闻
# ============================================================

async def fetch_news(url, source, max_items=20):
    async with httpx.AsyncClient(timeout=15, follow_redirects=True) as c:
        r = await c.get(url, headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code != 200:
            return []
        soup = BeautifulSoup(r.text, "xml")
        items = []
        for item in soup.find_all("item")[:max_items]:
            title = item.find("title")
            link = item.find("link")
            desc = item.find("description")
            if title:
                raw = desc.get_text(strip=True) if desc else ""
                clean = BeautifulSoup(raw, "html.parser").get_text(strip=True)[:200] if raw else ""
                items.append({"title": title.get_text(strip=True), "url": link.get_text(strip=True) if link else "", "source": source, "desc": clean})
        return items


async def fetch_kotaku_news():
    return await fetch_news("https://kotaku.com/rss", "Kotaku", 15)


async def fetch_pcgamer_news():
    return await fetch_news("https://www.pcgamer.com/rss/", "PCGamer", 15)


async def get_ign():
    for feed in ["https://www.ign.com/rss/articles", "https://feeds.feedburner.com/ign/games-all"]:
        result = await fetch_news(feed, "IGN")
        if result:
            return result
    return []


async def get_gamespot():
    return await fetch_news("https://www.gamespot.com/feeds/mashup/", "GameSpot")


# ============================================================
# 缓存
# ============================================================

_cache = {}
_content_refresh_count = 0

def cached(key, fn=None, ttl=600):
    now = datetime.now().timestamp()
    if key in _cache and now - _cache[key]["t"] < ttl:
        return _cache[key]["d"]
    return None


def cache_set(key, data):
    _cache[key] = {"d": data, "t": datetime.now().timestamp()}


# ============================================================
# 项目管理
# ============================================================

PROJ_FILE = BASE_DIR / "projects.json"

def load_proj():
    if PROJ_FILE.exists():
        return json.loads(PROJ_FILE.read_text("utf-8"))
    default = [
        {"id": 1, "name": "拳皇15", "stage": "运营中", "platform": "PS5/PC", "dev": "SNK", "dl": "已上线"},
        {"id": 2, "name": "双影奇境", "stage": "营销中", "platform": "PS5", "dev": "Ultizero", "dl": "2026-Q3"},
        {"id": 3, "name": "侍魂 晓", "stage": "运营中", "platform": "PS4/PS5", "dev": "SNK", "dl": "已上线"},
        {"id": 4, "name": "新项目A", "stage": "谈判中", "platform": "待定", "dev": "待确认", "dl": "待定"},
        {"id": 5, "name": "新项目B", "stage": "本地化中", "platform": "PC/PS5", "dev": "海外工作室", "dl": "2026-Q4"},
    ]
    save_proj(default)
    return default

def save_proj(data):
    PROJ_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), "utf-8")


# ============================================================
# 路由
# ============================================================

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/api/dashboard")
async def api_dashboard():
    d = cached("dash")
    if d:
        return d
    try:
        feat = await get_featured()
    except Exception as e:
        print(f"Featured error: {e}")
        feat = {"top_sellers": [], "new_releases": [], "specials": []}
    try:
        ign = await get_ign()
    except Exception as e:
        print(f"IGN error: {e}")
        ign = []
    try:
        gs = await get_gamespot()
    except Exception as e:
        print(f"GS error: {e}")
        gs = []
    try:
        kotaku = await fetch_kotaku_news()
    except Exception as e:
        print(f"Kotaku error: {e}")
        kotaku = []
    try:
        pcgamer = await fetch_pcgamer_news()
    except Exception as e:
        print(f"PCGamer error: {e}")
        pcgamer = []
    result = {**feat, "ign": ign, "gs": gs, "kotaku": kotaku, "pcgamer": pcgamer, "time": datetime.now().strftime("%Y-%m-%d %H:%M")}
    cache_set("dash", result)
    return result


@app.get("/api/search")
async def api_search(q: str):
    return await search_games(q)


@app.get("/api/game/{app_id}")
async def api_game(app_id: int):
    return await game_details(app_id)


@app.get("/api/reviews/{app_id}")
async def api_reviews(app_id: int, num: int = 20):
    return await game_reviews(app_id, num)


@app.get("/api/translate")
async def api_translate(text: str):
    return {"text": await translate(text)}


@app.get("/api/projects")
async def api_projects():
    return load_proj()


@app.post("/api/projects")
async def api_add_project(request: Request):
    body = await request.json()
    data = load_proj()
    body["id"] = max((p["id"] for p in data), default=0) + 1
    data.append(body)
    save_proj(data)
    return {"ok": True}


@app.put("/api/projects/{pid}")
async def api_update_project(pid: int, request: Request):
    body = await request.json()
    data = load_proj()
    for p in data:
        if p["id"] == pid:
            p.update(body)
            break
    save_proj(data)
    return {"ok": True}


@app.delete("/api/projects/{pid}")
async def api_delete_project(pid: int):
    data = [p for p in load_proj() if p["id"] != pid]
    save_proj(data)
    return {"ok": True}


@app.get("/api/compare")
async def api_compare(ids: str):
    id_list = [int(x) for x in ids.split(",") if x.strip().isdigit()]
    results = []
    # 并发获取
    async def fetch_one(aid):
        det = await game_details(aid)
        rev = await game_reviews(aid, num=5)
        if det:
            total = rev["positive"] + rev["negative"]
            return {
                **det,
                "positive": rev["positive"],
                "negative": rev["negative"],
                "rate": round(rev["positive"] / total * 100, 1) if total else 0,
                "score": rev["score"],
            }
        return None
    tasks = [fetch_one(aid) for aid in id_list]
    results = await asyncio.gather(*tasks)
    return [r for r in results if r]


@app.get("/api/report/{app_id}")
async def api_report(app_id: int):
    det = await game_details(app_id)
    rev = await game_reviews(app_id, num=20)
    if not det:
        return {"error": "not found"}
    total = rev["positive"] + rev["negative"]
    return {
        "game": det,
        "stats": {
            "positive": rev["positive"],
            "negative": rev["negative"],
            "rate": round(rev["positive"] / total * 100, 1) if total else 0,
            "score": rev["score"],
            "total": total,
        },
        "reviews": rev["reviews"],
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


@app.get("/api/report/{app_id}/docx")
async def api_report_docx(app_id: int):
    """生成docx格式的舆情报表"""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    det = await game_details(app_id)
    rev = await game_reviews(app_id, num=30)
    if not det:
        return {"error": "not found"}

    total = rev["positive"] + rev["negative"]
    rate = round(rev["positive"] / total * 100, 1) if total else 0

    doc = Document()
    title = doc.add_heading(f'{det["name"]} - 舆情分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('一、游戏基本信息', level=1)
    table = doc.add_table(rows=6, cols=2, style='Table Grid')
    for i, (k, v) in enumerate([
        ("游戏名称", det["name"]),
        ("开发商", ", ".join(det["developers"])),
        ("游戏类型", ", ".join(det["genres"])),
        ("发售日期", det["release_date"]),
        ("Steam链接", det["url"]),
        ("报告生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)

    doc.add_heading('二、评价数据概览', level=1)
    table2 = doc.add_table(rows=5, cols=2, style='Table Grid')
    for i, (k, v) in enumerate([
        ("总评价数", f"{total:,}"),
        ("好评数", f"{rev['positive']:,}"),
        ("差评数", f"{rev['negative']:,}"),
        ("好评率", f"{rate}%"),
        ("评价等级", rev["score"]),
    ]):
        table2.rows[i].cells[0].text = k
        table2.rows[i].cells[1].text = v

    doc.add_heading('三、评价分布分析', level=1)
    if rate >= 80:
        analysis = f"该游戏好评率达到{rate}%，属于高口碑游戏。玩家整体满意度较高，游戏品质得到市场认可。"
    elif rate >= 60:
        analysis = f"该游戏好评率为{rate}%，口碑中等偏上。存在一定比例的负面反馈，建议关注玩家提出的具体问题。"
    else:
        analysis = f"该游戏好评率仅为{rate}%，口碑较差。需要重点关注玩家反馈的负面问题，及时进行优化改进。"
    doc.add_paragraph(analysis)

    doc.add_heading('四、近期玩家评价样本', level=1)
    doc.add_paragraph(f"以下为最近{min(15, len(rev['reviews']))}条玩家评价：")
    for rv in rev["reviews"][:15]:
        tag = "好评" if rv["up"] else "差评"
        text = rv["text"][:500] if rv["text"] else "（评价内容为空）"
        p = doc.add_paragraph()
        run = p.add_run(f"[{tag}] ")
        run.bold = True
        run.font.color.rgb = RGBColor(34, 197, 94) if rv["up"] else RGBColor(239, 68, 68)
        p.add_run(f"时长: {rv['hours']}h | {rv['lang']}\n{text}\n")

    doc.add_heading('五、总结与建议', level=1)
    if rate >= 80:
        doc.add_paragraph("1. 游戏品质优秀，市场口碑良好，适合加大推广力度")
        doc.add_paragraph("2. 可重点关注玩家提到的亮点功能，在营销中突出宣传")
        doc.add_paragraph("3. 持续监控负面评价，及时响应玩家诉求")
    elif rate >= 60:
        doc.add_paragraph("1. 游戏整体表现中等，建议针对性优化玩家反馈集中的问题")
        doc.add_paragraph("2. 分析差评原因，制定改进计划")
        doc.add_paragraph("3. 加强社区运营，提升玩家满意度")
    else:
        doc.add_paragraph("1. 游戏口碑较差，建议暂停推广，优先解决核心问题")
        doc.add_paragraph("2. 深入分析差评原因，制定全面优化方案")
        doc.add_paragraph("3. 考虑延期发布，确保游戏品质达标")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"report_{app_id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ============================================================
# Content Copilot
# ============================================================

PLATFORM_TONES = {
    "bilibili": {"name": "B站", "hook": "深度看点", "style": "适合做3-5分钟解析"},
    "douyin": {"name": "抖音", "hook": "强钩子", "style": "适合做15-30秒快节奏视频"},
    "xiaohongshu": {"name": "小红书", "hook": "种草/避坑", "style": "适合做清单式体验分享"},
    "wechat": {"name": "公众号", "hook": "行业观察", "style": "适合做日报或周报条目"},
}

DEMO_BLOCK_TERMS = [
    "adult", "hentai", "sex", "erotic", "nude", "nsfw", "porn",
    "waifu", "dating", "empress", "undercover train", "strip",
    "casino", "gamble", "gambling", "poker", "slot", "gnome",
    "succubus", "tentacle", "maid", "bikini", "ntr",
]

CURATED_CONTENT_OPPORTUNITIES = [
    {
        "id": "curated-phantom-blade-zero",
        "app_id": None,
        "type": "game",
        "source": "内容精选",
        "bucket": "curated",
        "title": "影之刃零 Phantom Blade Zero",
        "summary": "国产动作游戏具备强视觉、强战斗卖点，适合做海外玩家反应、实机解析和主机玩家向内容。",
        "angle": "国产动作游戏如何打动海外主机玩家？",
        "image": "",
        "url": "",
        "heat": 96,
        "tags": ["国产动作", "主机发行", "海外关注"],
        "recommended_formats": ["B站解析", "30秒脚本", "海外资讯", "采访提纲"],
    },
    {
        "id": "curated-black-myth",
        "app_id": None,
        "type": "game",
        "source": "内容精选",
        "bucket": "curated",
        "title": "黑神话：悟空 内容长尾",
        "summary": "高认知国产IP适合做长尾内容复盘，包括海外评价、玩家讨论、主机表现和同类游戏机会。",
        "angle": "爆款之后，国产游戏还能继续讲什么内容？",
        "image": "",
        "url": "",
        "heat": 94,
        "tags": ["国产IP", "海外口碑", "长尾运营"],
        "recommended_formats": ["公众号复盘", "B站解析", "小红书笔记", "日报条目"],
    },
    {
        "id": "curated-elden-ring-nightreign",
        "app_id": None,
        "type": "game",
        "source": "内容精选",
        "bucket": "curated",
        "title": "艾尔登法环：黑夜君临",
        "summary": "魂系IP自带讨论度，适合做难度争议、联机体验、购买建议和玩家情绪分析。",
        "angle": "魂系玩家为什么会关注这次联机玩法？",
        "image": "",
        "url": "",
        "heat": 93,
        "tags": ["魂系", "玩家争议", "购买建议"],
        "recommended_formats": ["15秒短视频", "30秒脚本", "B站解析", "风险提示"],
    },
    {
        "id": "curated-split-fiction",
        "app_id": None,
        "type": "game",
        "source": "内容精选",
        "bucket": "curated",
        "title": "双影奇境 Split Fiction",
        "summary": "合作叙事游戏适合做双人体验、情侣/朋友游玩场景、玩法亮点和发行节奏内容。",
        "angle": "合作游戏为什么适合短视频平台传播？",
        "image": "",
        "url": "",
        "heat": 91,
        "tags": ["合作游戏", "短视频传播", "发行项目"],
        "recommended_formats": ["抖音标题", "小红书笔记", "30秒脚本", "采访提纲"],
    },
    {
        "id": "curated-steam-next-fest",
        "app_id": None,
        "type": "news",
        "source": "内容精选",
        "bucket": "curated",
        "title": "Steam Next Fest 试玩节内容机会",
        "summary": "试玩节适合提前建立选题池和素材表，活动开始后快速产出试玩清单、黑马预测和海外反馈。",
        "angle": "发行团队如何从试玩节里发现潜力项目？",
        "image": "",
        "url": "",
        "heat": 90,
        "tags": ["活动规划", "试玩节", "选题池"],
        "recommended_formats": ["活动前准备", "日报条目", "B站清单", "采访问题"],
    },
    {
        "id": "curated-gamescom",
        "app_id": None,
        "type": "news",
        "source": "内容精选",
        "bucket": "curated",
        "title": "Gamescom 展会跟拍准备",
        "summary": "线下展会需要提前规划必拍镜头、采访对象、B-roll素材和活动后复盘，适合展示视频岗价值。",
        "angle": "游戏展会跟拍到底应该提前准备什么？",
        "image": "",
        "url": "",
        "heat": 88,
        "tags": ["展会跟拍", "采访", "素材清单"],
        "recommended_formats": ["拍摄清单", "采访提纲", "30秒脚本", "内部汇报"],
    },
]


def demo_svg_image(label, bg="#e0f2fe", fg="#2563eb"):
    safe_label = re.sub(r"[<>&]", "", label)
    svg = (
        "<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 640 360'>"
        f"<rect width='640' height='360' fill='{bg}'/>"
        "<circle cx='560' cy='80' r='140' fill='rgba(255,255,255,.34)'/>"
        "<circle cx='80' cy='320' r='180' fill='rgba(255,255,255,.28)'/>"
        f"<text x='48' y='178' font-size='42' font-weight='700' font-family='Arial, sans-serif' fill='{fg}'>{safe_label}</text>"
        f"<text x='50' y='226' font-size='20' font-family='Arial, sans-serif' fill='{fg}' opacity='.72'>Content Opportunity</text>"
        "</svg>"
    )
    return "data:image/svg+xml;charset=utf-8," + svg.replace("#", "%23").replace(" ", "%20")


EVENT_ONLY_OPPORTUNITY_IDS = {"curated-gamescom", "rank-tga-watch"}


DEMO_IMAGE_BY_ID = {
    "curated-phantom-blade-zero": demo_svg_image("Phantom Blade Zero", "#e0e7ff", "#3730a3"),
    "curated-black-myth": "https://cdn.akamai.steamstatic.com/steam/apps/2358720/header.jpg",
    "curated-elden-ring-nightreign": "https://cdn.akamai.steamstatic.com/steam/apps/2622380/header.jpg",
    "curated-split-fiction": "https://cdn.akamai.steamstatic.com/steam/apps/2001120/header.jpg",
    "curated-steam-next-fest": demo_svg_image("Steam Next Fest"),
}


DEMO_RANK_OPPORTUNITIES = [
    {
        "id": "rank-monster-hunter-wilds",
        "type": "game",
        "source": "Steam热度榜",
        "bucket": "rank",
        "title": "Monster Hunter Wilds",
        "summary": "大型动作共斗游戏适合做版本更新、玩家回流、装备流派和多人联机体验内容。",
        "angle": "共斗游戏更新后，玩家最关心哪些变化？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/2246340/header.jpg",
        "url": "https://store.steampowered.com/app/2246340/",
        "heat": 87,
        "trend": "+18%",
        "tags": ["Steam热度", "共斗", "版本更新"],
        "recommended_formats": ["30秒脚本", "B站解析", "小红书清单", "日报条目"],
    },
    {
        "id": "rank-stellar-blade",
        "type": "game",
        "source": "Steam热度榜",
        "bucket": "rank",
        "title": "Stellar Blade",
        "summary": "主机移植与PC玩家讨论度高，适合做配置表现、画面对比、入坑建议和海外评价整理。",
        "angle": "主机游戏登陆PC后，内容团队应该先看什么？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/3489700/header.jpg",
        "url": "https://store.steampowered.com/app/3489700/",
        "heat": 86,
        "trend": "+15%",
        "tags": ["主机移植", "PC表现", "海外口碑"],
        "recommended_formats": ["短视频切片", "购买建议", "海外资讯", "封面文案"],
    },
    {
        "id": "rank-marvel-rivals",
        "type": "game",
        "source": "Steam热度榜",
        "bucket": "rank",
        "title": "Marvel Rivals",
        "summary": "多人竞技游戏适合持续跟进赛季更新、角色强度、玩家情绪和社区争议。",
        "angle": "竞技游戏赛季内容怎样拆成连续选题？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/2767030/header.jpg",
        "url": "https://store.steampowered.com/app/2767030/",
        "heat": 85,
        "trend": "+11%",
        "tags": ["赛季更新", "竞技", "玩家情绪"],
        "recommended_formats": ["15秒短视频", "30秒脚本", "争议点分析", "日报"],
    },
    {
        "id": "rank-silk-song",
        "type": "game",
        "source": "愿望单观察",
        "bucket": "rank",
        "title": "Hollow Knight: Silksong",
        "summary": "高期待独立游戏适合做愿望单观察、发布节奏、玩家期待和同类游戏推荐。",
        "angle": "高期待独立游戏为什么适合提前做内容池？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/1030300/header.jpg",
        "url": "https://store.steampowered.com/app/1030300/",
        "heat": 84,
        "trend": "+9%",
        "tags": ["独立游戏", "愿望单", "期待值"],
        "recommended_formats": ["B站解析", "小红书笔记", "选题池", "日报"],
    },
    {
        "id": "rank-hades-ii",
        "type": "game",
        "source": "Steam热度榜",
        "bucket": "rank",
        "title": "Hades II",
        "summary": "高口碑续作适合做版本更新、流派推荐、玩家评价和同类动作Roguelike内容。",
        "angle": "高口碑续作更新后，内容团队怎么判断是否值得跟进？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/1145350/header.jpg",
        "url": "https://store.steampowered.com/app/1145350/",
        "heat": 83,
        "trend": "+8%",
        "tags": ["Roguelike", "口碑", "版本更新"],
        "recommended_formats": ["B站解析", "短视频脚本", "小红书清单", "日报"],
    },
    {
        "id": "rank-dont-starve-together",
        "type": "game",
        "source": "Steam热度榜",
        "bucket": "rank",
        "title": "Don't Starve Together",
        "summary": "长线运营游戏适合做促销节点、好友联机、入坑指南和老玩家回流内容。",
        "angle": "长线联机游戏为什么适合反复做入坑内容？",
        "image": "https://cdn.akamai.steamstatic.com/steam/apps/322330/header.jpg",
        "url": "https://store.steampowered.com/app/322330/",
        "heat": 81,
        "trend": "+6%",
        "tags": ["长线运营", "联机", "入坑指南"],
        "recommended_formats": ["入坑建议", "30秒脚本", "促销提醒", "日报条目"],
    },
    {
        "id": "rank-indie-radar",
        "type": "news",
        "source": "海外资讯",
        "bucket": "rank",
        "title": "独立游戏黑马监控",
        "summary": "适合从Steam新品、试玩节和Reddit讨论中筛选潜力项目，形成每日选题池。",
        "angle": "发行团队怎样更早发现可能爆的独立游戏？",
        "image": demo_svg_image("Indie Radar", "#dcfce7", "#15803d"),
        "url": "",
        "heat": 80,
        "trend": "+7%",
        "tags": ["黑马预测", "试玩节", "选题池"],
        "recommended_formats": ["日报条目", "B站清单", "短视频脚本", "市场观察"],
    },
]


def normalize_demo_opportunities(opportunities, limit=18):
    cleaned = [item for item in opportunities if item.get("id") not in EVENT_ONLY_OPPORTUNITY_IDS]
    seen = {item.get("id") for item in cleaned}
    for item in DEMO_RANK_OPPORTUNITIES:
        if item["id"] not in seen:
            cleaned.append(item.copy())
            seen.add(item["id"])
    for idx, item in enumerate(cleaned):
        item.setdefault("trend", f"+{max(5, 26 - idx * 2)}%")
        if not item.get("image"):
            item["image"] = DEMO_IMAGE_BY_ID.get(item.get("id"), demo_svg_image("Game Radar"))
    cleaned.sort(key=lambda item: item.get("heat", 0), reverse=True)
    return cleaned[:limit]


def is_demo_safe(title):
    low = (title or "").lower()
    return not any(term in low for term in DEMO_BLOCK_TERMS)


def text_score(text):
    score = 50
    hot_words = ["launch", "release", "trailer", "update", "review", "steam", "ps5", "xbox", "switch", "sales"]
    risk_words = ["delay", "bug", "controversy", "refund", "layoff", "lawsuit", "server", "review bomb"]
    low = text.lower()
    score += sum(6 for w in hot_words if w in low)
    score += sum(5 for w in risk_words if w in low)
    return min(score, 96)


def content_tags(name, source="", desc=""):
    raw = f"{name} {source} {desc}".lower()
    tags = []
    if any(w in raw for w in ["sale", "discount", "special"]):
        tags.append("促销节点")
    if any(w in raw for w in ["review", "评价", "score"]):
        tags.append("口碑观察")
    if any(w in raw for w in ["trailer", "showcase", "direct", "展会"]):
        tags.append("视频素材")
    if any(w in raw for w in ["steam", "热销", "top"]):
        tags.append("Steam热度")
    if not tags:
        tags = ["内容机会", "海外资讯"]
    return tags[:3]


def opportunity_from_game(item, bucket, rank):
    name = item.get("name") or "未知游戏"
    discount = item.get("discount_percent", 0)
    if bucket == "top_sellers":
        why = "Steam热销榜靠前，说明玩家购买意愿正在增强，适合当天快速做热点跟进。"
        angle = "这款游戏为什么突然被玩家买爆？"
        urgency = 92 - rank * 3
    elif bucket == "new_releases":
        why = "新品进入Steam曝光位，适合做首发信息整合、玩法看点和购买建议。"
        angle = "新游上线值不值得关注？"
        urgency = 84 - rank * 2
    else:
        why = f"当前有{discount}%折扣，适合做促销提醒、入坑建议和同类游戏对比。"
        angle = "打折后现在是不是入坑好时机？"
        urgency = 78 - rank * 2 + min(discount, 50) // 5
    return {
        "id": f"steam-{bucket}-{item.get('id')}",
        "app_id": item.get("id"),
        "type": "game",
        "source": "Steam",
        "bucket": bucket,
        "title": name,
        "summary": why,
        "angle": angle,
        "image": item.get("header_image", ""),
        "url": item.get("url", ""),
        "heat": max(45, min(98, urgency)),
        "tags": content_tags(name, "steam"),
        "recommended_formats": ["15秒短视频", "30秒脚本", "B站解析", "日报条目"],
    }


def opportunity_from_news(item, rank):
    title = item.get("title", "")
    desc = item.get("desc", "")
    src = item.get("source", "News")
    score = text_score(f"{title} {desc}") - rank
    return {
        "id": f"news-{src}-{rank}",
        "app_id": None,
        "type": "news",
        "source": src,
        "bucket": "news",
        "title": title,
        "summary": desc or "海外媒体出现新动态，适合整理成中文资讯和内容选题。",
        "angle": "这条海外资讯对国内玩家/发行团队意味着什么？",
        "image": "",
        "url": item.get("url", ""),
        "heat": max(40, min(95, score)),
        "tags": content_tags(title, src, desc),
        "recommended_formats": ["海外资讯", "公众号日报", "短视频口播"],
    }


def content_profile(item):
    title = item.get("title", "今日热点")
    tags = item.get("tags", []) or []
    summary = item.get("summary", "")
    angle = item.get("angle", "为什么值得关注")
    text = f"{title} {summary} {' '.join(tags)} {angle}".lower()

    if any(k in text for k in ["展会", "gamescom", "采访", "活动", "试玩节"]):
        return {
            "audience": "游戏发行/内容团队、关注活动资讯的核心玩家",
            "selling_point": "活动前准备、现场素材和活动后复盘",
            "emotion": "专业、信息密度高、像团队内部作战清单",
            "proof": "活动节点、采访对象、必拍镜头、海外反馈",
            "cta": "收藏清单，活动当天直接照着拍",
        }
    if any(k in text for k in ["魂系", "难度", "艾尔登", "黑夜"]):
        return {
            "audience": "魂系玩家、观望购买的动作游戏玩家",
            "selling_point": "难度争议、联机体验和入坑判断",
            "emotion": "克制但有冲突感，先抛问题再给判断",
            "proof": "玩家讨论、玩法变化、购买建议、风险点",
            "cta": "评论区选：实机解析、购买建议还是避坑",
        }
    if any(k in text for k in ["国产", "主机", "海外", "影之刃", "黑神话"]):
        return {
            "audience": "主机玩家、国产游戏关注者、海外发行团队",
            "selling_point": "国产游戏的海外传播点和主机玩家接受度",
            "emotion": "自信、专业，避免空喊国产情怀",
            "proof": "视觉表现、战斗卖点、海外玩家反应、平台调性",
            "cta": "想看海外反应还是实机拆解，评论区选",
        }
    if any(k in text for k in ["合作", "双人", "split", "朋友", "情侣"]):
        return {
            "audience": "双人游戏玩家、情侣/朋友开黑用户、轻度内容平台用户",
            "selling_point": "双人场景、情绪传播和短视频切片",
            "emotion": "轻快、强场景感、适合种草",
            "proof": "合作机制、冲突笑点、关卡变化、玩家反应",
            "cta": "转给你的搭子，看看谁更适合玩",
        }
    return {
        "audience": "正在寻找新游戏和热点资讯的玩家",
        "selling_point": "玩家关注点、内容传播点和购买判断",
        "emotion": "直接、清楚、有判断",
        "proof": "Steam热度、玩家讨论、媒体信息、玩法卖点",
        "cta": "想看详细评测/购买建议/同类推荐，评论区选一个",
    }


def variant_pick(options, variant=0):
    if not options:
        return ""
    return options[variant % len(options)]


def build_titles(title, angle, item=None, variant=0):
    clean = re.sub(r"\s+", " ", title).strip()
    profile = content_profile(item or {"title": title, "angle": angle})
    hook_a = variant_pick(["真正值得看的是这3点", "别只看热度，先看内容价值", "玩家会不会买账，关键在这里"], variant)
    hook_b = variant_pick(["30秒讲清", "一条视频讲明白", "先别急着跟风"], variant + 1)
    return {
        "bilibili": [
            f"{clean}为什么值得做一期？{hook_a}",
            f"从{clean}看{profile['selling_point']}：内容团队该怎么切",
            f"{clean}：亮点、风险和玩家关注点一次讲清",
        ],
        "douyin": [
            f"{hook_b}：{clean}为什么突然值得关注？",
            f"{clean}别乱跟，这3个点最适合做内容",
            f"如果今天只做一条游戏热点，我会先看{clean}",
        ],
        "xiaohongshu": [
            f"{clean}适合现在入坑/关注吗？先看这份笔记",
            f"{clean}内容选题笔记：受众、亮点、风险",
            f"刷到{clean}别急着收藏，先看它适合谁",
        ],
        "wechat": [
            f"内容观察：{clean}带来的选题与传播机会",
            f"从{clean}看近期玩家关注趋势与内容切口",
            f"{clean}热点观察：{angle}",
        ],
    }


def build_script(title, angle, item=None, seconds=30, variant=0):
    profile = content_profile(item or {"title": title, "angle": angle})
    source = (item or {}).get("source", "内容源")
    tags = " / ".join((item or {}).get("tags", [])[:3]) or profile["proof"]
    opening = variant_pick([
        f"今天这个热点不是单纯火了，关键是它给内容团队留了一个切口：{angle}",
        f"如果今天只能选一个游戏热点跟进，我会先看{title}，原因有三个。",
        f"{title}现在适不适合做内容？不要只看热度，要看玩家到底在讨论什么。",
    ], variant)
    if seconds <= 15:
        beats = [
            f"0-3秒｜口播：{title}今天值得关注，但不是因为它名字大。",
            f"3-10秒｜画面：放商店页/预告片/评论截图。口播：核心看点是“{angle}”，适合切{profile['selling_point']}。",
            f"10-15秒｜口播：后续想看实机、评测还是避坑？评论区选一个。",
        ]
    elif seconds <= 30:
        beats = [
            f"0-3秒｜钩子：{opening}",
            f"3-8秒｜背景：来源来自{source}，标签是{tags}。先给观众一个判断：它适合做{profile['selling_point']}。",
            f"8-18秒｜三点拆解：第一，看玩家为什么关注；第二，看有没有可剪的画面/争议点；第三，看它能不能转成购买建议或资讯日报。",
            f"18-25秒｜画面建议：切商店页、预告片高能段、玩家评论/媒体标题，用字幕标出“亮点/风险/适合谁”。",
            f"25-30秒｜收口：{profile['cta']}。",
        ]
    else:
        beats = [
            f"开场｜今天的重点是{title}。它不是单纯一条新闻，而是一个可以拆成多平台内容的选题。",
            f"第一段｜为什么值得做：围绕“{angle}”，先解释玩家关注的原因，再给出内容判断。",
            f"第二段｜玩家关注点：受众是{profile['audience']}，核心卖点是{profile['selling_point']}，证据可以从{profile['proof']}里找。",
            "第三段｜剪辑结构：先用强画面做钩子，再放来源截图建立可信度，中段拆亮点和风险，最后给观众一个明确选择。",
            "第四段｜平台拆分：B站做解析，抖音做30秒钩子，小红书做入坑/避坑，公众号做日报条目。",
            f"结尾｜互动：{profile['cta']}。",
        ]
    return "\n".join(f"{i + 1}. {beat}" for i, beat in enumerate(beats))


def build_work_order_markdown(title, angle, item, profile, pack):
    source_url = item.get("url")
    lines = [
        f"# {title} 剪辑交付工作单",
        "",
        "## 1. 今天为什么剪",
        f"- 推荐角度：{angle}",
        f"- 目标受众：{profile['audience']}",
        f"- 核心价值：{profile['selling_point']}",
        "- 判断标准：能不能在30秒内讲清“为什么值得关注、适合谁、风险是什么”。",
        "",
        "## 2. 30秒成片结构",
        "- 0-3秒：强问题开场，必须出现游戏名和核心判断。",
        "- 3-8秒：放来源画面，证明不是空泛跟热点。",
        "- 8-18秒：三段信息，分别讲亮点、玩家关注点、争议/风险。",
        "- 18-25秒：给购买建议、入坑建议或内容判断。",
        "- 25-30秒：评论区问题，引导用户选择下一条内容方向。",
        "",
        "## 3. 素材准备",
        "- 01_key_visual：主视觉 / Steam页 / 官方图。",
        "- 02_gameplay：实机或预告片高能片段3段。",
        "- 03_social_proof：玩家评论、海外媒体标题、热度截图。",
        "- 04_compare：同类游戏或平台标签对比素材。",
        "- 05_cover：封面底图、标题短句、字幕关键词。",
        "",
        "## 4. 剪辑要求",
        "- 画幅：9:16优先，兼容16:9二次裁切。",
        "- 字幕：每屏不超过18个字，只保留关键词。",
        "- 节奏：前10秒至少3次画面变化，避免纯口播。",
        "- 封面：一版问题式，一版判断式，一版入坑/避坑式。",
        "",
        "## 5. 发布前核查",
        "- 游戏名、平台、发售/活动时间是否准确。",
        "- 素材是否来自官方或可引用来源。",
        "- 标题是否过度承诺或容易引战。",
        "- 评论截图是否遮挡个人隐私信息。",
        "",
        "## 6. 团队分工",
        "- 剪辑：产出30秒短视频、字幕版和工程文件。",
        "- 运营：选平台标题，安排发布时间。",
        "- 素材：补齐来源链接和素材文件夹。",
        "- 负责人：审核风险点，决定是否扩展成长视频。",
        "",
        "## 7. 可直接复制的首版标题",
        f"- B站：{pack['titles']['bilibili'][0]}",
        f"- 抖音：{pack['titles']['douyin'][0]}",
        f"- 小红书：{pack['titles']['xiaohongshu'][0]}",
    ]
    if source_url:
        lines.extend(["", "## 8. 来源", f"- {source_url}"])
    return "\n".join(lines)


def build_content_pack(item):
    title = item.get("title", "今日热点")
    angle = item.get("angle", "为什么值得关注")
    variant = int(item.get("_variant", 0) or 0)
    profile = content_profile(item)
    titles = build_titles(title, angle, item, variant)
    material_source = "Steam页面 / 官方预告片 / 玩家评论截图" if item.get("type") == "game" else "海外原文链接 / 媒体标题截图 / 社媒讨论截图"
    shot_list = [
        {"time": "0-3秒", "screen": "最强画面或标题截图", "voice": variant_pick([f"{title}今天值得关注，但别只看热度。", f"{title}现在最适合切的不是新闻，而是玩家讨论。"], variant), "edit": "大字标题 + 快速推近"},
        {"time": "3-8秒", "screen": "来源页、商店页或媒体页", "voice": f"它的内容切口是：{angle}", "edit": "标注来源，建立可信度"},
        {"time": "8-18秒", "screen": "三段素材：玩法/评论/平台信息", "voice": f"适合抓三个点：{profile['selling_point']}、{profile['proof']}、潜在风险。", "edit": "三连分屏或节奏切点"},
        {"time": "18-25秒", "screen": "玩家评论或同类游戏对比", "voice": f"如果面向{profile['audience']}，标题要更具体，别只写'又火了'。", "edit": "字幕强调受众"},
        {"time": "25-30秒", "screen": "封面文案/评论区引导", "voice": profile["cta"], "edit": "停留1秒，方便截图转发"},
    ]
    owner_tasks = [
        {"role": "剪辑", "task": "按30秒分镜剪一条竖版短视频，保留工程和字幕文件", "output": "30秒短视频1条 + 字幕版"},
        {"role": "运营", "task": "从标题池选择平台标题，按平台调整语气并安排发布时间", "output": "B站/抖音/小红书发布文案"},
        {"role": "素材", "task": "补齐官方页面、预告片、评论截图和授权来源", "output": "素材文件夹 + 来源链接"},
        {"role": "负责人", "task": "检查事实、标题尺度、素材版权和风险点", "output": "发布确认"},
    ]
    pack = {
        "source": item,
        "analysis": {
            "why_hot": item.get("summary", ""),
            "player_focus": [
                profile["selling_point"],
                profile["proof"],
                "受众是否清晰",
                "是否能转成短视频钩子或日报条目",
            ],
            "content_opportunity": f"围绕“{angle}”做一条短视频，再沉淀成日报条目。",
            "risk": "当前为自动整理结果，发布前建议人工确认事实、发售日期和素材版权。",
        },
        "titles": titles,
        "scripts": {
            "15s": build_script(title, angle, item, 15, variant),
            "30s": build_script(title, angle, item, 30, variant),
            "60s": build_script(title, angle, item, 60, variant),
        },
        "shot_list": shot_list,
        "cover_copy": [
            f"{title}",
            variant_pick(["真正值得看的3个点", "别只看热度", "内容团队先看这里"], variant),
            variant_pick(["亮点 / 风险 / 受众", "适合谁？怎么剪？", "30秒选题判断"], variant + 1),
        ],
        "publish_copy": {
            "douyin": f"{title}不是简单跟热点，关键是“{angle}”。这条先拆玩家为什么关注、内容怎么切、发布前要确认什么。#游戏资讯 #Steam #游戏推荐",
            "xiaohongshu": f"今天的游戏内容选题：{title}\n\n适合人群：{profile['audience']}\n内容角度：{angle}\n素材建议：{material_source}\n发布前记得确认来源和素材授权。",
            "bilibili_desc": f"本期从{title}切入，整理玩家关注点、内容传播角度和剪辑素材建议。核心问题：{angle}",
        },
        "asset_checklist": [
            material_source,
            "官方截图或预告片高能片段 3-5 个",
            "玩家评论/媒体标题截图 2-3 张",
            "同类游戏或平台标签对比素材",
            "封面主视觉、标题短句、字幕关键词",
        ],
        "risk_checklist": [
            "发售日期、平台、价格信息是否准确",
            "素材是否来自官方页面或可引用来源",
            "标题是否过度夸张或引战",
            "评论截图是否避免暴露个人隐私",
        ],
        "editing_notes": [
            "前3秒必须出现问题或判断，不要从背景慢慢讲。",
            "每7-8秒切一次视觉信息：画面、评论、标题、标签。",
            "字幕只保留关键词，避免整段口播全贴屏幕。",
            "封面不要写满，保留一个强问题和一个明确对象。",
        ],
        "interview": [
            f"这次{title}最希望玩家注意到的核心体验是什么？",
            "如果只能用一个镜头介绍这款游戏，你会选哪一段？",
            "团队如何看待当前玩家讨论中的争议点？",
            "后续版本或发行计划中，最值得期待的内容是什么？",
        ],
        "owner_tasks": owner_tasks,
        "value_summary": {
            "time_saved": "约90-150分钟",
            "team_value": "把选题判断、脚本、分镜、素材清单、发布文案和风险核查集中成一张工作单。",
            "best_use": "适合晨会后直接分配给剪辑、运营和素材整理同事。",
        },
    }
    pack["work_order_markdown"] = build_work_order_markdown(title, angle, item, profile, pack)
    return pack


def merge_curated_opportunities(opportunities, min_count=18, rotation=0):
    seen = {item["id"] for item in opportunities}
    title_seen = set()
    merged = []
    curated = [item.copy() for item in CURATED_CONTENT_OPPORTUNITIES]
    if curated:
        offset = rotation % len(curated)
        curated = curated[offset:] + curated[:offset]
    for item in curated:
        title_key = item["title"].strip().lower()
        if item["id"] not in seen and title_key not in title_seen:
            merged.append(item.copy())
            title_seen.add(title_key)
            seen.add(item["id"])
    for item in opportunities:
        title_key = item["title"].strip().lower()
        if title_key in title_seen:
            continue
        merged.append(item)
        title_seen.add(title_key)
    if len(merged) < min_count:
        return merged
    return merged[:max(min_count, len(CURATED_CONTENT_OPPORTUNITIES))]


def fallback_topic_opportunity(topic):
    clean = re.sub(r"\s+", " ", topic).strip()
    return {
        "id": f"topic-{int(datetime.now().timestamp())}",
        "app_id": None,
        "type": "custom",
        "source": "手动输入",
        "bucket": "topic",
        "title": clean,
        "summary": "手动输入的内容方向，适合快速生成平台标题、短视频脚本、封面文案和采访提纲。",
        "angle": f"{clean}为什么值得玩家关注？",
        "image": "",
        "url": "",
        "heat": 80,
        "tags": ["手动选题", "内容生成"],
        "recommended_formats": ["15秒短视频", "30秒脚本", "采访提纲", "日报条目"],
    }


async def build_topic_opportunity(topic):
    clean = re.sub(r"\s+", " ", topic).strip()
    if not clean:
        return fallback_topic_opportunity("今日游戏热点")

    for item in CURATED_CONTENT_OPPORTUNITIES:
        if clean.lower() in item["title"].lower() or item["title"].lower() in clean.lower():
            result = item.copy()
            result["source"] = "精选池匹配"
            return result

    try:
        matches = await search_games(clean)
    except Exception as e:
        print(f"Topic search error: {e}")
        return fallback_topic_opportunity(clean)

    if not matches:
        return fallback_topic_opportunity(clean)

    first = matches[0]
    app_id = first.get("app_id")
    try:
        det, rev = await asyncio.gather(game_details(app_id), game_reviews(app_id, num=5))
    except Exception as e:
        print(f"Topic detail error: {e}")
        det, rev = None, {"positive": 0, "negative": 0, "score": "", "reviews": []}

    title = (det or {}).get("name") or first.get("name") or clean
    total = rev.get("positive", 0) + rev.get("negative", 0)
    rate = round(rev.get("positive", 0) / total * 100, 1) if total else 0
    genres = (det or {}).get("genres", [])[:3]
    genre_text = "、".join(genres) if genres else "玩法、口碑和购买决策"
    score_text = f"当前Steam好评率约{rate}%，评价标签为{rev.get('score')}" if total else "当前适合先从玩法卖点、玩家讨论和内容传播点切入"

    return {
        "id": f"topic-steam-{app_id}",
        "app_id": app_id,
        "type": "game",
        "source": "Steam搜索",
        "bucket": "topic",
        "title": title,
        "summary": f"{title}与{genre_text}相关，{score_text}，适合生成购买建议、玩法亮点和平台选题。",
        "angle": f"{title}现在最适合从哪个角度做内容？",
        "image": (det or {}).get("header_image") or first.get("image", ""),
        "url": (det or {}).get("url") or f"https://store.steampowered.com/app/{app_id}",
        "heat": 86 if total else 78,
        "tags": content_tags(title, "Steam搜索", " ".join(genres)),
        "recommended_formats": ["15秒短视频", "30秒脚本", "B站解析", "购买建议"],
        "review_stats": {
            "positive": rev.get("positive", 0),
            "negative": rev.get("negative", 0),
            "rate": rate,
            "score": rev.get("score", ""),
        },
    }


async def get_content_opportunities(refresh=False):
    global _content_refresh_count
    if refresh:
        _content_refresh_count += 1
        _cache.pop("content_ops", None)
    cached_data = cached("content_ops")
    if cached_data:
        return cached_data

    try:
        dash = await asyncio.wait_for(api_dashboard(), timeout=8)
    except Exception as e:
        print(f"Content opportunities dashboard fallback: {e}")
        dash = {
            "top_sellers": [],
            "new_releases": [],
            "specials": [],
            "ign": [],
            "gs": [],
            "pcgamer": [],
            "kotaku": [],
        }
    opportunities = []
    for bucket in ["top_sellers", "new_releases", "specials"]:
        for rank, item in enumerate(dash.get(bucket, [])[:5], start=1):
            if not is_demo_safe(item.get("name", "")):
                continue
            opportunities.append(opportunity_from_game(item, bucket, rank))

    news_items = []
    for key in ["ign", "gs", "pcgamer", "kotaku"]:
        news_items.extend(dash.get(key, [])[:4])
    for rank, item in enumerate(news_items[:10], start=1):
        if not is_demo_safe(item.get("title", "")):
            continue
        opportunities.append(opportunity_from_news(item, rank))

    opportunities.sort(key=lambda x: x["heat"], reverse=True)
    opportunities = merge_curated_opportunities(opportunities, rotation=_content_refresh_count)
    opportunities = normalize_demo_opportunities(opportunities)
    result = {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "stats": {
            "opportunities": len(opportunities),
            "short_video_topics": min(28, len(opportunities) * 2),
            "news_items": len([o for o in opportunities if o.get("type") == "news"]),
            "risk_items": len([o for o in opportunities if o["heat"] >= 85]),
        },
        "opportunities": opportunities,
        "hero": opportunities[0] if opportunities else None,
        "refresh_id": _content_refresh_count,
    }
    cache_set("content_ops", result)
    return result


@app.get("/api/content/opportunities")
async def api_content_opportunities(refresh: int = 0):
    return await get_content_opportunities(refresh=bool(refresh))


@app.post("/api/content/generate")
async def api_content_generate(request: Request):
    item = await request.json()
    return build_content_pack(item)


@app.get("/api/content/topic")
async def api_content_topic(q: str):
    return await build_topic_opportunity(q)


@app.get("/api/content/daily")
async def api_content_daily():
    data = await get_content_opportunities()
    items = data["opportunities"][:8]
    lines = [
        f"# GamePoch Content Daily - {data['time']}",
        "",
        "## 今日重点",
    ]
    for i, item in enumerate(items[:3], start=1):
        lines.append(f"{i}. {item['title']} - {item['summary']}")
    lines.extend(["", "## 推荐选题"])
    for item in items[:5]:
        lines.append(f"- {item['title']}：{item['angle']}")
    lines.extend(["", "## 今日可执行动作"])
    lines.extend([
        "- 选1个Steam热度最高条目做15-30秒短视频。",
        "- 选2条海外资讯整理成中文日报。",
        "- 对高热度但存在争议的游戏，先做避坑/购买建议角度。",
    ])
    return {"time": data["time"], "markdown": "\n".join(lines), "items": items}


def pick_focus_events(limit=2):
    priority = {"高": 0, "中": 1, "低": 2}
    return sorted(EVENT_PLANNER_DATA, key=lambda e: (priority.get(e["priority"], 9), e["name"]))[:limit]


@app.get("/api/content/brief")
async def api_content_brief():
    data = await get_content_opportunities()
    items = data["opportunities"][:8]
    focus_events = pick_focus_events(2)
    time_text = datetime.now().strftime("%Y-%m-%d %H:%M")

    lines = [
        f"# GamePoch 内容作战简报 - {time_text}",
        "",
        "## 1. 今日结论",
        f"- 今日发现 {data['stats']['opportunities']} 个内容机会，建议优先产出 {min(3, len(items))} 条短视频和 1 份海外资讯整理。",
        f"- 高优先级风险 {data['stats']['risk_items']} 个，发布前需要人工确认事实、素材版权和标题尺度。",
        f"- 活动侧建议优先准备：{'、'.join(e['name'] for e in focus_events)}。",
        "",
        "## 2. 今日热点优先级",
    ]

    for i, item in enumerate(items[:5], start=1):
        lines.append(f"{i}. {item['title']}（{item['source']} / 热度 {item['heat']}）")
        lines.append(f"   - 内容角度：{item['angle']}")
        lines.append(f"   - 建议产出：{', '.join(item['recommended_formats'][:3])}")

    lines.extend(["", "## 3. 推荐执行动作"])
    actions = [
        "剪辑：先做1条15-30秒热点短视频，用强钩子测试反馈。",
        "运营：把前5个热点整理成平台选题池，分配到B站/抖音/小红书。",
        "英语资讯：选择2条海外资讯做中文摘要，保留原文链接方便复核。",
        "负责人：对高热度但可能有争议的内容先审标题，再安排发布。",
    ]
    for action in actions:
        lines.append(f"- {action}")

    lines.extend(["", "## 4. 活动内容准备"])
    for event in focus_events:
        lines.append(f"### {event['name']}")
        lines.append(f"- 时间：{event['date']}")
        lines.append(f"- 类型：{event['type']}；优先级：{event['priority']}")
        lines.append(f"- 活动前：{event['pre_event']['topics'][0]}")
        lines.append(f"- 当天必拍：{', '.join(event['live_event']['must_shoot'][:4])}")
        lines.append(f"- 活动后：{event['post_event']['recap_daily']}")

    lines.extend(["", "## 5. 风险提醒"])
    risks = []
    for item in items[:5]:
        risks.append(f"{item['title']}：自动生成角度需人工确认事实和素材来源。")
    for event in focus_events:
        risks.extend(f"{event['name']}：{risk}" for risk in event["risks"][:1])
    for risk in risks[:6]:
        lines.append(f"- {risk}")

    lines.extend(["", "## 6. 可直接复制给团队的任务分工"])
    lines.extend([
        "- 视频：今天先剪热点短视频1条，保留素材链接和封面文案。",
        "- 运营：把标题池同步到发布表，标注平台、角度和发布时间。",
        "- 英语：完成海外资讯摘要，标注来源、时间和可引用信息。",
        "- 活动：根据活动规划清单补齐采访问题和拍摄物料。",
    ])

    return {
        "time": time_text,
        "markdown": "\n".join(lines),
        "hotspots": items,
        "events": [event_summary(e) for e in focus_events],
        "actions": actions,
    }


@app.get("/api/content/brief/docx")
async def api_content_brief_docx():
    """导出正式Word版内容作战简报。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    brief = await api_content_brief()
    doc = Document()

    title = doc.add_heading("GamePoch 内容作战简报", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"生成时间：{brief['time']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in brief["markdown"].splitlines()[2:]:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        elif line.startswith("   - "):
            p = doc.add_paragraph(style="List Bullet 2")
            p.add_run(line[5:])
        else:
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("附录：使用建议", level=1)
    tips = [
        "团队晨会优先查看今日工作台，再进入内容生成和活动规划。",
        "如果临时需要跟进某款游戏，可在内容生成区直接输入游戏名。",
        "简报可作为负责人日报、周报或PPT素材底稿。",
    ]
    for tip in tips:
        doc.add_paragraph(tip, style="List Bullet")

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            if paragraph.style.name.startswith("Heading"):
                run.font.color.rgb = RGBColor(15, 118, 110)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=GamePoch_Content_Brief.docx"},
    )


@app.get("/api/content/calendar")
async def api_content_calendar():
    data = await get_content_opportunities()
    opportunities = data["opportunities"][:10]
    if not opportunities:
        opportunities = merge_curated_opportunities([])[:10]
    if not opportunities:
        return {
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "days": [],
            "markdown": "# GamePoch 7日内容发布排期\n\n暂无可用内容机会，请稍后刷新。",
            "summary": {"days": 0, "tasks": 0, "platforms": []},
        }
    focus_events = pick_focus_events(2)
    start = datetime.now()
    platforms = ["B站", "抖音", "小红书", "公众号"]
    formats = ["30秒短视频", "热点解析", "图文笔记", "日报条目"]
    owners = ["视频剪辑", "内容运营", "英语资讯", "活动跟拍"]
    calendar = []

    for day_index in range(7):
        date_text = (start).strftime("%Y-%m-%d")
        day_label = ["今天", "明天", "第3天", "第4天", "第5天", "第6天", "第7天"][day_index]
        tasks = []
        for slot in range(2):
            item = opportunities[(day_index * 2 + slot) % len(opportunities)]
            platform = platforms[(day_index + slot) % len(platforms)]
            content_format = formats[(day_index + slot) % len(formats)]
            tasks.append({
                "time": "上午" if slot == 0 else "下午",
                "platform": platform,
                "format": content_format,
                "topic": item["title"],
                "angle": item["angle"],
                "owner": owners[(day_index + slot) % len(owners)],
                "asset": "Steam页面/官方素材/评论截图" if item["type"] == "game" else "海外原文链接/中文摘要/来源截图",
                "goal": "测试热点反馈" if slot == 0 else "沉淀可复用素材",
            })
        if day_index in [1, 4] and focus_events:
            event = focus_events[0 if day_index == 1 else min(1, len(focus_events) - 1)]
            tasks.append({
                "time": "收尾",
                "platform": "团队内部",
                "format": "活动准备",
                "topic": event["name"],
                "angle": event["content_opportunities"][0],
                "owner": "活动跟拍",
                "asset": "采访问题/必拍镜头/B-roll清单",
                "goal": "提前准备活动内容素材",
            })
        calendar.append({
            "date": date_text,
            "label": day_label,
            "tasks": tasks,
        })
        start = start + timedelta(days=1)

    markdown = [f"# GamePoch 7日内容发布排期 - {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
    for day in calendar:
        markdown.append(f"## {day['label']}｜{day['date']}")
        for task in day["tasks"]:
            markdown.append(f"- {task['time']}｜{task['platform']}｜{task['format']}｜{task['topic']}")
            markdown.append(f"  - 角度：{task['angle']}")
            markdown.append(f"  - 负责人：{task['owner']}；素材：{task['asset']}；目标：{task['goal']}")
        markdown.append("")

    return {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "days": calendar,
        "markdown": "\n".join(markdown),
        "summary": {
            "days": len(calendar),
            "tasks": sum(len(day["tasks"]) for day in calendar),
            "platforms": platforms,
        },
    }


# ============================================================
# Game Event Content Planner
# ============================================================

EVENT_PLANNER_DATA = [
    {
        "id": "steam-next-fest",
        "name": "Steam Next Fest",
        "date": "2026-06-16 至 2026-06-23（规划样例）",
        "type": "线上试玩节",
        "status": "活动前准备",
        "priority": "高",
        "focus_genres": ["独立游戏", "动作冒险", "策略模拟", "合作游戏"],
        "content_opportunities": [
            "提前筛选值得试玩的新游，做“编辑部试玩清单”。",
            "对有潜力的Demo做15秒首印象短视频。",
            "整理海外玩家对试玩版的第一波反馈。"
        ],
        "risks": [
            "试玩版质量波动大，发布推荐前需要人工体验确认。",
            "素材使用需确认官方媒体包或商店页授权范围。"
        ],
        "pre_event": {
            "topics": [
                "Steam新品试玩节最值得关注的10款游戏",
                "这几款Demo可能会成为下一个独立游戏黑马",
                "发行团队如何从试玩节里判断潜力项目"
            ],
            "titles": {
                "bilibili": "Steam新品试玩节前瞻：这10款Demo值得提前加入愿望单",
                "douyin": "试玩节别乱点，先看这5款最有爆相的新游",
                "xiaohongshu": "Steam Next Fest试玩清单：适合周末体验的新游"
            },
            "overseas_summary": "海外玩家通常会集中讨论玩法新鲜度、Demo完成度、优化表现和愿望单转化。内容团队可提前建立试玩记录表，活动开始后快速更新。",
            "interview_questions": [
                "这次Demo最希望玩家验证哪一项核心体验？",
                "团队会如何根据试玩节反馈调整后续开发？",
                "如果用一句话推荐这款游戏，你希望玩家记住什么？"
            ],
            "materials": ["Steam商店截图", "官方预告片", "开发者Press Kit", "试玩录屏模板", "中文标题与标签表"]
        },
        "live_event": {
            "must_shoot": ["开场倒计时/活动页面", "愿望单按钮", "Demo实机第一分钟", "亮点玩法片段", "玩家评论截图"],
            "interview_targets": ["制作人", "发行负责人", "现场试玩玩家", "海外媒体编辑"],
            "b_roll": ["操作手部特写", "游戏Logo与主视觉", "展台/直播页面环境", "玩家排队或试玩反应"],
            "short_clips": ["15秒试玩第一印象", "30秒玩法亮点", "玩家一句话评价", "避坑提醒"]
        },
        "post_event": {
            "script_30s": "0-3秒：这次Steam试玩节最值得看的不是数量，而是谁真的有爆相。3-12秒：我们筛了玩法完成度、玩家反馈和传播点。12-24秒：重点推荐三类：合作玩法、强视觉独立游戏、容易做短视频切片的动作游戏。24-30秒：完整清单已经整理好，评论区告诉我们你想先看哪款。",
            "recap_daily": "今日重点：完成试玩节候选清单初筛，优先跟进高讨论度Demo。建议动作：剪3条短视频、整理1份海外玩家反馈、向团队同步潜力项目。",
            "social_copy": "试玩节不是随便下载Demo，而是内容团队发现黑马的窗口。我们整理了几款值得提前关注的新游，从玩法亮点、传播点和风险三个维度快速判断是否值得跟进。",
            "internal_summary": "本次活动适合沉淀为选题库、试玩素材库和潜力项目观察表。建议活动后48小时内完成首轮内容发布。"
        }
    },
    {
        "id": "gamescom",
        "name": "Gamescom",
        "date": "2026-08-19 至 2026-08-23（规划样例）",
        "type": "线下展会 / 海外媒体活动",
        "status": "物料规划",
        "priority": "高",
        "focus_genres": ["主机游戏", "动作RPG", "射击游戏", "多人联机"],
        "content_opportunities": [
            "适合做现场跟拍、展台探访和制作人采访。",
            "可把海外媒体试玩反馈整理成中文内容。",
            "适合为合作伙伴输出活动复盘和传播建议。"
        ],
        "risks": [
            "现场噪音和人流会影响采访收音。",
            "部分试玩区域可能禁止拍摄，需要提前确认规则。"
        ],
        "pre_event": {
            "topics": ["Gamescom前瞻：今年海外玩家最期待什么", "展会跟拍必备镜头清单", "制作人采访前必须准备的8个问题"],
            "titles": {
                "bilibili": "Gamescom前瞻：这些游戏可能会成为展会最大话题",
                "douyin": "去Gamescom跟拍，最不能漏的5个镜头",
                "xiaohongshu": "游戏展会拍摄清单：新手也能拍出专业素材"
            },
            "overseas_summary": "海外媒体关注新预告、试玩体验、开发者访谈和平台独占消息。内容团队应提前准备关键词、游戏名单和采访对象优先级。",
            "interview_questions": ["这次展会最核心的展示目标是什么？", "玩家试玩后最常问的问题是什么？", "海外发行最担心的市场反馈是什么？"],
            "materials": ["无线麦克风", "备用电池", "采访授权确认", "展台路线图", "中英双语问题卡"]
        },
        "live_event": {
            "must_shoot": ["展馆外景", "展台全景", "试玩队伍", "制作人介绍", "玩家试玩反应", "周边/海报/主视觉"],
            "interview_targets": ["制作人", "市场负责人", "KOL", "海外媒体", "试玩玩家"],
            "b_roll": ["Logo墙", "手柄操作", "屏幕实机", "展台互动", "媒体证件和活动物料"],
            "short_clips": ["展台30秒巡礼", "制作人一句话推荐", "玩家试玩反馈", "现场最火游戏排行"]
        },
        "post_event": {
            "script_30s": "0-3秒：今年Gamescom现场，玩家最关注的其实是这几类游戏。3-12秒：我们看了试玩排队、展台热度和媒体反馈。12-24秒：动作RPG、多人联机和强视觉游戏更容易形成传播。24-30秒：后续我们会整理完整采访和试玩观察。",
            "recap_daily": "今日重点：完成展会素材采集与采访记录。建议动作：当天先发现场短视频，24小时内发布采访切片，48小时内产出复盘长文。",
            "social_copy": "展会内容不是只拍热闹，关键是拍到可复用素材：玩家反应、制作人观点、实机亮点和现场传播氛围。",
            "internal_summary": "Gamescom适合建立海外内容素材库。建议按游戏、采访对象、镜头类型和授权状态整理素材。"
        }
    },
    {
        "id": "tga",
        "name": "The Game Awards",
        "date": "2026-12-10（规划样例）",
        "type": "颁奖典礼 / 新作发布",
        "status": "选题储备",
        "priority": "中",
        "focus_genres": ["年度游戏", "3A新作", "独立游戏", "预告片"],
        "content_opportunities": [
            "适合做提名前瞻、获奖预测和预告片速递。",
            "适合把海外讨论整理成中文热点日报。",
            "活动后可快速产出获奖复盘和玩家争议点。"
        ],
        "risks": ["直播时间可能不适合国内团队值守。", "获奖争议容易引战，标题需避免过度绝对化。"],
        "pre_event": {
            "topics": ["TGA获奖预测：哪些游戏最有冠军相", "今年TGA最值得等的新预告", "年度游戏争议点提前看"],
            "titles": {
                "bilibili": "TGA前瞻：今年年度游戏悬念到底在哪里？",
                "douyin": "TGA别只看获奖，这些新预告更重要",
                "xiaohongshu": "TGA观看清单：适合游戏玩家提前收藏"
            },
            "overseas_summary": "海外讨论集中在年度游戏归属、首发预告、嘉宾阵容和平台发布节奏。内容团队可提前准备候选游戏资料卡。",
            "interview_questions": ["你认为今年获奖结果会影响后续销售吗？", "玩家争议点集中在哪些方面？", "哪个新预告最可能带来愿望单增长？"],
            "materials": ["提名名单", "候选游戏截图", "直播时间表", "快速字幕模板", "获奖结果记录表"]
        },
        "live_event": {
            "must_shoot": ["获奖瞬间", "新预告首帧", "观众弹幕/评论", "社媒热词截图"],
            "interview_targets": ["媒体编辑", "核心玩家", "发行市场同事"],
            "b_roll": ["奖杯画面", "候选游戏混剪", "直播界面", "社媒趋势榜"],
            "short_clips": ["获奖30秒速报", "新预告亮点", "争议结果复盘", "玩家反应合集"]
        },
        "post_event": {
            "script_30s": "0-3秒：今年TGA结果出来后，争议最大的不是谁获奖。3-12秒：真正值得关注的是新预告和玩家讨论方向。12-24秒：这些内容会影响接下来几周的搜索、愿望单和平台讨论。24-30秒：我们整理了最值得跟进的三个选题。",
            "recap_daily": "今日重点：完成TGA获奖与预告片信息整理。建议动作：先发速报，再做争议点解析，最后沉淀年度趋势复盘。",
            "social_copy": "TGA对内容团队的价值，不只在奖项，而在它集中释放了玩家下一阶段会讨论的话题。",
            "internal_summary": "TGA适合做热点追踪和年度内容复盘，建议提前安排夜间值守和次日剪辑排期。"
        }
    },
    {
        "id": "summer-game-fest",
        "name": "Summer Game Fest",
        "date": "2026-06-05（规划样例）",
        "type": "线上发布会 / 预告片集中发布",
        "status": "复盘沉淀",
        "priority": "中",
        "focus_genres": ["预告片", "动作游戏", "独立新作", "发行节点"],
        "content_opportunities": ["适合做预告片速看。", "适合按品类整理重点游戏。", "适合为发行团队判断海外关注趋势。"],
        "risks": ["预告片信息密集，容易遗漏来源。", "部分内容只有英文信息，需人工复核翻译。"],
        "pre_event": {
            "topics": ["发布会前应该关注哪些厂商", "预告片速看模板提前准备", "海外资讯整理值班表"],
            "titles": {
                "bilibili": "Summer Game Fest前瞻：今年哪些新作值得等？",
                "douyin": "今晚发布会，看这几个重点就够了",
                "xiaohongshu": "游戏发布会观看攻略：别错过这些新消息"
            },
            "overseas_summary": "发布会期间信息密度高，应提前准备游戏名、厂商、平台、发售窗口、素材链接五列记录表。",
            "interview_questions": ["这次发布节奏为什么选择这个时间点？", "预告片最想传递的核心卖点是什么？", "目标玩家是谁？"],
            "materials": ["直播链接", "记录表", "预告片下载链接", "字幕模板", "封面模板"]
        },
        "live_event": {
            "must_shoot": ["直播开场", "重点预告片", "发售日画面", "平台信息", "社媒热评"],
            "interview_targets": ["内容编辑", "核心玩家", "市场同事"],
            "b_roll": ["预告片关键帧", "直播画面", "社媒趋势截图", "游戏Logo"],
            "short_clips": ["一分钟发布会速看", "最值得等的三款", "玩家反应", "发售日汇总"]
        },
        "post_event": {
            "script_30s": "0-3秒：这场发布会信息很多，但真正值得记的是这三件事。3-12秒：第一是重点新作，第二是发售窗口，第三是玩家讨论最热的预告。12-24秒：内容团队可以立刻拆成速报、解析和清单。24-30秒：完整汇总已经整理好。",
            "recap_daily": "今日重点：完成发布会重点游戏与发售窗口整理。建议动作：当天发速报，次日做重点游戏解析。",
            "social_copy": "发布会内容多不等于都值得做。真正该抓的是玩家会搜索、会争论、会收藏的信息。",
            "internal_summary": "发布会适合建立快速资讯机制，建议固定模板减少整理时间。"
        }
    },
    {
        "id": "state-of-play",
        "name": "PlayStation State of Play",
        "date": "待官方公布（规划样例）",
        "type": "平台发布会",
        "status": "持续关注",
        "priority": "中",
        "focus_genres": ["PS5", "主机独占", "日系游戏", "动作冒险"],
        "content_opportunities": ["适合关注主机平台发行节奏。", "适合整理索尼生态合作机会。", "适合输出给合作伙伴的中文摘要。"],
        "risks": ["官方公布前信息不确定，避免使用未经确认爆料。", "平台独占信息需准确引用。"],
        "pre_event": {
            "topics": ["下一场State of Play可能看什么", "PS5玩家最期待的新消息", "发行团队如何准备平台发布会内容"],
            "titles": {
                "bilibili": "State of Play前瞻：PS玩家最该关注哪些消息？",
                "douyin": "索尼发布会前，先看这几个可能重点",
                "xiaohongshu": "PS5新作消息整理：发布会前收藏这一篇"
            },
            "overseas_summary": "平台发布会通常围绕独占内容、第三方合作、DLC和发售日期。应提前建立游戏名单和平台标签。",
            "interview_questions": ["这款游戏与PlayStation玩家的匹配点是什么？", "平台合作对发行节奏有什么影响？", "素材发布后最希望玩家讨论什么？"],
            "materials": ["PlayStation官方链接", "候选游戏清单", "平台标签", "索尼合作方资料", "中文摘要模板"]
        },
        "live_event": {
            "must_shoot": ["开场主视觉", "平台Logo", "新作画面", "发售日期", "独占信息"],
            "interview_targets": ["主机玩家", "市场负责人", "内容编辑"],
            "b_roll": ["PS5主机/手柄", "游戏实机片段", "官方海报", "社媒讨论"],
            "short_clips": ["PS发布会重点速看", "玩家最期待", "独占消息整理", "一分钟平台观察"]
        },
        "post_event": {
            "script_30s": "0-3秒：这次State of Play最值得看的不是数量，而是平台信号。3-12秒：哪些游戏获得重点曝光，哪些发售窗口被确认。12-24秒：这些信息会影响主机玩家的购买和关注。24-30秒：我们整理了适合继续跟进的内容方向。",
            "recap_daily": "今日重点：跟进PlayStation发布会重点游戏。建议动作：做中文速报、平台合作观察和玩家评论整理。",
            "social_copy": "平台发布会最适合看发行节奏。哪些游戏被重点展示，往往意味着接下来几周的内容机会。",
            "internal_summary": "State of Play适合服务主机发行和合作伙伴汇报，建议保留准确来源链接。"
        }
    },
    {
        "id": "nintendo-direct",
        "name": "Nintendo Direct",
        "date": "待官方公布（规划样例）",
        "type": "平台发布会",
        "status": "持续关注",
        "priority": "中",
        "focus_genres": ["Switch", "合家欢", "独立游戏", "任天堂IP"],
        "content_opportunities": ["适合做Switch玩家资讯整理。", "适合关注独立游戏曝光机会。", "适合做轻量化短视频清单。"],
        "risks": ["任天堂信息节奏严格，避免提前传播未确认消息。", "IP素材版权使用需谨慎。"],
        "pre_event": {
            "topics": ["Nintendo Direct前瞻：Switch玩家该期待什么", "任天堂发布会最适合做的短视频选题", "独立游戏如何借平台发布会获得关注"],
            "titles": {
                "bilibili": "Nintendo Direct前瞻：哪些消息最值得Switch玩家等？",
                "douyin": "任天堂发布会，重点看这几类游戏",
                "xiaohongshu": "Switch玩家发布会观看清单"
            },
            "overseas_summary": "任天堂发布会信息适合做轻量化整理，重点关注首发游戏、DLC、独立游戏合集和发售日期。",
            "interview_questions": ["这款游戏适合Switch玩家的原因是什么？", "如果面向家庭用户，最该突出哪个卖点？", "发布会曝光后如何承接玩家关注？"],
            "materials": ["任天堂官方链接", "Switch游戏清单", "发售日期表", "轻量封面模板", "中文标签"]
        },
        "live_event": {
            "must_shoot": ["发布会标题页", "游戏Logo", "发售日期", "多人/家庭玩法", "玩家评论"],
            "interview_targets": ["Switch玩家", "家庭用户", "内容编辑"],
            "b_roll": ["Switch掌机画面", "多人游玩场景", "游戏主视觉", "社媒趋势"],
            "short_clips": ["Switch新作一分钟", "合家欢游戏推荐", "独立游戏亮点", "发售日汇总"]
        },
        "post_event": {
            "script_30s": "0-3秒：这次Nintendo Direct，Switch玩家最该记住这几条。3-12秒：先看新作，再看发售日，最后看适合多人玩的内容。12-24秒：这些游戏更适合做清单、种草和家庭娱乐角度。24-30秒：完整整理已经准备好。",
            "recap_daily": "今日重点：整理Nintendo Direct重点内容。建议动作：发布轻量清单、发售日期图文和Switch玩家向短视频。",
            "social_copy": "任天堂发布会适合做轻松、清单式内容。玩家更关心什么时候玩、和谁玩、值不值得买。",
            "internal_summary": "Nintendo Direct适合沉淀Switch内容选题库，注意素材版权和官方来源引用。"
        }
    }
]


def event_summary(event):
    return {
        "id": event["id"],
        "name": event["name"],
        "date": event["date"],
        "type": event["type"],
        "status": event["status"],
        "priority": event["priority"],
        "focus_genres": event["focus_genres"],
        "content_opportunities": event["content_opportunities"],
        "risks": event["risks"],
    }


@app.get("/api/events/planner")
async def api_event_planner():
    return {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "stats": {
            "events": len(EVENT_PLANNER_DATA),
            "high_priority": len([e for e in EVENT_PLANNER_DATA if e["priority"] == "高"]),
            "shooting_tasks": sum(len(e["live_event"]["must_shoot"]) for e in EVENT_PLANNER_DATA),
            "content_packages": len(EVENT_PLANNER_DATA) * 3,
        },
        "events": [event_summary(e) for e in EVENT_PLANNER_DATA],
    }


@app.get("/api/events/planner/{event_id}")
async def api_event_planner_detail(event_id: str):
    for event in EVENT_PLANNER_DATA:
        if event["id"] == event_id:
            return event
    return {"error": "event not found"}


@app.get("/api/demo/status")
async def api_demo_status():
    """演示前自检：确认核心模块都有可展示内容。"""
    checks = []

    try:
        ops = await get_content_opportunities()
        checks.append({
            "name": "热点发现",
            "ok": bool(ops.get("hero")) and len(ops.get("opportunities", [])) >= 6,
            "detail": f"{len(ops.get('opportunities', []))} 个内容机会",
        })
    except Exception as exc:
        checks.append({"name": "热点发现", "ok": False, "detail": str(exc)})

    try:
        topic = await build_topic_opportunity("Elden Ring")
        pack = build_content_pack(topic)
        checks.append({
            "name": "自定义内容生成",
            "ok": bool(topic.get("title")) and bool(pack.get("titles")) and bool(pack.get("scripts", {}).get("30s")),
            "detail": f"{topic.get('title', '未生成标题')} / {len(pack.get('titles', []))} 个平台标题",
        })
    except Exception as exc:
        checks.append({"name": "自定义内容生成", "ok": False, "detail": str(exc)})

    try:
        events = await api_event_planner()
        checks.append({
            "name": "活动规划",
            "ok": len(events.get("events", [])) >= 6,
            "detail": f"{len(events.get('events', []))} 个活动内容包",
        })
    except Exception as exc:
        checks.append({"name": "活动规划", "ok": False, "detail": str(exc)})

    try:
        calendar = await api_content_calendar()
        checks.append({
            "name": "7日发布排期",
            "ok": calendar.get("summary", {}).get("days") == 7 and calendar.get("summary", {}).get("tasks", 0) >= 14,
            "detail": f"{calendar.get('summary', {}).get('days', 0)} 天 / {calendar.get('summary', {}).get('tasks', 0)} 条任务",
        })
    except Exception as exc:
        checks.append({"name": "7日发布排期", "ok": False, "detail": str(exc)})

    try:
        brief = await api_content_brief()
        checks.append({
            "name": "简报导出",
            "ok": "内容作战简报" in brief.get("markdown", "") and len(brief.get("hotspots", [])) > 0,
            "detail": f"{len(brief.get('hotspots', []))} 个热点进入简报",
        })
    except Exception as exc:
        checks.append({"name": "简报导出", "ok": False, "detail": str(exc)})

    failed = [check for check in checks if not check["ok"]]
    return {
        "ready": len(failed) == 0,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "title": "GamePoch Content Copilot 系统自检",
        "positioning": "游戏发行团队内容效率工作台",
        "checks": checks,
        "next_step": "可以使用：按今日工作台、热点发现、内容生成、活动规划、发布排期、简报导出的顺序完成内容工作。"
        if not failed else "需要先修复未通过项，再继续使用。",
    }


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
